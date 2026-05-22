#!/usr/bin/env python3
"""
md_to_docx.py - convert a legal memo markdown draft to a docx.

Used by legal-memo-writer plugin at the export phase. Invoked from the main
session via Bash. Reads markdown and applies the visual spec from
lib/docx-render/README.md (Arial 12pt body, 11pt italic blockquote,
1-inch margins, single line spacing + 6pt after, justified, bold-paragraph
headings without Word Heading styles, blockquote left-indent).

Optionally prepends a yellow warning banner for forced-exit or
manual-review memos.

Requires python-docx (`pip install python-docx`).
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Cm, Inches, Pt, RGBColor
    from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is not installed. Install with `pip install python-docx`.\n"
        "Fallback: try `pandoc <input> -o <output>` or surface the markdown path.\n"
    )
    sys.exit(2)


# ---------------------------------------------------------------------------
# Visual spec (from lib/docx-render/README.md, ported from the user's
# Cowork org-level skill `legal-memo-style 11.skill`).
# ---------------------------------------------------------------------------

FONT_NAME = "Arial"
FONT_SIZE_BODY = Pt(12)
FONT_SIZE_QUOTE = Pt(11)
MARGIN = Inches(1)                  # 1440 twips, 2.54 cm
INDENT_FIRSTLINE_BODY = Cm(1.11)    # 630 DXA
INDENT_LEFT_BLOCKQUOTE = Cm(1.59)   # 900 DXA
INDENT_LEFT_LIST = Cm(1.27)         # 720 DXA = 0.5" — bullets indented deeper
                                    # than body first-line so the marker sits
                                    # at ~360 DXA (left - inherited hanging=360)
                                    # and the wrapped text starts at 720 DXA.
                                    # Matches the canonical Cowork visual spec
                                    # (manual Word UI override on the source
                                    # docx: <w:ind w:left="720"/> + tab clear
                                    # at 360 on top of the ListBullet style).
SPACING_BEFORE = Pt(0)
SPACING_AFTER = Pt(6)               # 120 DXA = 6pt after
LINE_SPACING_VALUE = 1.0            # single

WARNING_BG = "FFF3CD"
WARNING_BORDER = "FFE69C"


# ---------------------------------------------------------------------------
# Low-level XML helpers (cell shading + borders, used by the banner)
# ---------------------------------------------------------------------------

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:color"), hex_color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


# ---------------------------------------------------------------------------
# Document-level setup
# ---------------------------------------------------------------------------

def apply_page_setup(doc):
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN


def configure_default_style(doc):
    """Set the Normal style to Arial 12pt with single spacing + 6pt after.

    All paragraphs created later override font and indent explicitly per
    paragraph type — but setting Normal correctly is the safety net for any
    unstyled run.
    """
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY
    # East-Asian font fallback so Cyrillic stays in Arial in some Word builds
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING_VALUE
    pf.space_before = SPACING_BEFORE
    pf.space_after = SPACING_AFTER


# ---------------------------------------------------------------------------
# Paragraph builders — one helper per paragraph type from the spec table
# ---------------------------------------------------------------------------

def _apply_std_paragraph_format(p, first_line_indent=None, left_indent=None):
    """Apply the standard spacing + justification to a paragraph."""
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = LINE_SPACING_VALUE
    pf.space_before = SPACING_BEFORE
    pf.space_after = SPACING_AFTER
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent


def _style_run(run, *, bold=False, italic=False, size=None):
    run.font.name = FONT_NAME
    run.font.size = size if size is not None else FONT_SIZE_BODY
    run.bold = bold
    run.italic = italic
    # Ensure east-asian font also applied to the run (otherwise Cyrillic
    # can fall back to default Calibri/Times on some Word builds).
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)


def add_heading_paragraph(doc, text):
    """Bold body-size paragraph, no indent. Used for ALL heading levels.

    Per user's spec, we deliberately do NOT use Word's Heading 1..4 styles.
    The numbering ('1.', '1.1.') comes from the markdown text itself, which
    keeps numbering stable across Word versions.
    """
    p = doc.add_paragraph()
    _apply_std_paragraph_format(p, first_line_indent=Cm(0))
    render_inline(p, text, bold_outer=True)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    _apply_std_paragraph_format(p, first_line_indent=INDENT_FIRSTLINE_BODY)
    render_inline(p, text)


def add_blockquote_paragraph(doc, text):
    p = doc.add_paragraph()
    _apply_std_paragraph_format(
        p,
        first_line_indent=Cm(0),
        left_indent=INDENT_LEFT_BLOCKQUOTE,
    )
    render_inline(p, text, italic_outer=True, size=FONT_SIZE_QUOTE)


def _clear_inherited_tab(paragraph, pos_dxa):
    """Insert <w:tabs><w:tab w:val="clear" w:pos="N"/></w:tabs> into pPr.

    Used after a list paragraph has its left indent overridden, to clear the
    tab stop inherited from the ListBullet / ListNumber style at the default
    marker position. Without the clear, some Word versions tab the wrapped
    line back to the old position instead of honoring the new w:ind.
    """
    pPr = paragraph._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "clear")
    tab.set(qn("w:pos"), str(pos_dxa))
    tabs.append(tab)
    # OOXML schema: w:tabs must appear before w:ind / w:jc inside w:pPr.
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        ind.addprevious(tabs)
    else:
        pPr.append(tabs)


def _disable_contextual_spacing(paragraph):
    """Insert <w:contextualSpacing w:val="0"/> into pPr.

    Word's built-in ListBullet / ListNumber styles set contextualSpacing=true
    in styles.xml, which suppresses the paragraph after-spacing between
    consecutive list items of the same style (the "Don't add space between
    paragraphs of the same style" checkbox in the Paragraph dialog). We
    override at paragraph level to restore the 6pt after-spacing between
    items, matching the canonical Cowork visual spec.
    """
    pPr = paragraph._element.get_or_add_pPr()
    cs = OxmlElement("w:contextualSpacing")
    cs.set(qn("w:val"), "0")
    # OOXML schema: contextualSpacing comes after w:ind and before w:jc.
    jc = pPr.find(qn("w:jc"))
    if jc is not None:
        jc.addprevious(cs)
    else:
        pPr.append(cs)


def add_list_item(doc, text, *, ordered):
    style_name = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style_name)
    # Both bullets and numbered lists get the same paragraph-level overrides:
    # left=720 DXA (deeper than body first-line indent), tab clear at 360 so
    # wrapped text snaps to the new indent, justified alignment, and
    # contextualSpacing=false so the 6pt after-spacing applies between items.
    # first_line_indent is NOT set: the hanging=360 inherited from the
    # ListBullet / ListNumber style stays in effect, so the marker (• or 1.)
    # sits at left-hanging = 360 DXA and wrapped text starts at 720 DXA.
    _apply_std_paragraph_format(p, left_indent=INDENT_LEFT_LIST)
    _clear_inherited_tab(p, 360)
    _disable_contextual_spacing(p)
    render_inline(p, text)


# ---------------------------------------------------------------------------
# Warning banner (for non-approved memos)
# ---------------------------------------------------------------------------

def add_warning_banner(doc, final_status, remaining_issues, fallback_banners=None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, WARNING_BG)
    set_cell_border(cell, WARNING_BORDER)

    title_p = cell.paragraphs[0]
    _apply_std_paragraph_format(title_p, first_line_indent=Cm(0))
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = "MANUAL REVIEW REQUIRED"
    subtitle_text = (
        "Manual check recommended before relying on this memorandum. "
        f"Final status: {final_status}."
    )
    if final_status and final_status.startswith("forced_exit"):
        title = "REVIEWER NOTES NOT FULLY RESOLVED"
    elif final_status and final_status.startswith("accepted_early"):
        title = "USER ACCEPTED EARLY — REMAINING ISSUES"
    elif final_status == "fallback_research_summary_delivered":
        # Phase 8 branch A — user picked research-summary mode via heartbeat.
        # Memo intentionally skipped IRAC analysis (no Risk, no Recommendation),
        # but the document IS a deliverable.
        title = "RESEARCH SUMMARY MODE — IRAC ANALYSIS NOT PERFORMED"
    elif final_status == "fallback_summary_delivered":
        # Universal catastrophic fallback per always-deliver.md. In practice this
        # path writes fallback-summary.md (not memo-<slug>.docx) and so doesn't
        # normally invoke md_to_docx.py — but if a future code path does call
        # the renderer with this status, the banner accurately labels it.
        title = "PIPELINE FALLBACK — RESEARCH INCOMPLETE"
    elif final_status and final_status.startswith("approved") and fallback_banners:
        # Approved memo with pipeline fallbacks (e.g. MCP rate-limited, partial
        # coverage). The reviewer loop approved the content; the banner is a
        # disclosure about *how* the research was conducted, not a warning that
        # the analysis is incomplete. Use a softer, accurate title.
        title = "PIPELINE FALLBACK NOTICE — REVIEW BEFORE CLIENT USE"
        subtitle_text = (
            "The memo was approved by the reviewer loop. The fallbacks listed below "
            "applied during research — verify items tagged in research files (e.g. "
            "[rate-limited fallback]) before client delivery. "
            f"Final status: {final_status}."
        )
    title_run = title_p.add_run(title)
    _style_run(title_run, bold=True)

    subtitle_p = cell.add_paragraph()
    _apply_std_paragraph_format(subtitle_p, first_line_indent=Cm(0))
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_p.add_run(subtitle_text)
    _style_run(subtitle_run)

    # Fallback banners: rendered above the per-issue list because they often
    # describe pipeline-level downgrades (research-summary, MCP unavailable,
    # rate-limited fallback) that contextualize WHY there are remaining issues.
    if fallback_banners:
        heading_p = cell.add_paragraph()
        _apply_std_paragraph_format(heading_p, first_line_indent=Cm(0))
        heading_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_run = heading_p.add_run("Pipeline fallbacks that fired during this run:")
        _style_run(h_run, bold=True)
        for banner in fallback_banners:
            li = cell.add_paragraph(style="List Bullet")
            _apply_std_paragraph_format(li, first_line_indent=Cm(0))
            li.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = li.add_run(str(banner))
            _style_run(run)

    if remaining_issues:
        heading_p = cell.add_paragraph()
        _apply_std_paragraph_format(heading_p, first_line_indent=Cm(0))
        heading_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_run = heading_p.add_run("Remaining blocking issues:")
        _style_run(h_run, bold=True)
        for issue in remaining_issues:
            li = cell.add_paragraph(style="List Bullet")
            _apply_std_paragraph_format(li, first_line_indent=Cm(0))
            li.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = li.add_run(issue)
            _style_run(run)

    doc.add_paragraph()  # spacer after banner


def extract_fallback_banners(state_path):
    """Read state.json.fallback_banners[]. Returns [] when state is unreachable
    or the field is empty/missing. Each banner is a free-text string written by
    a fallback path in always-deliver.md (heartbeat research-summary, MCP
    unavailable, MCP rate-limited, drafting-failed-to-summary, etc.).
    """
    if not state_path or not Path(state_path).exists():
        return []
    try:
        state = json.loads(Path(state_path).read_text(encoding="utf-8-sig"))
    except (json.JSONDecodeError, OSError):
        return []
    banners = state.get("fallback_banners")
    if not isinstance(banners, list):
        return []
    # Coerce each to string and drop empties/non-strings defensively.
    return [str(b) for b in banners if isinstance(b, (str, int, float)) and str(b).strip()]


def summarize_issue(issue):
    if isinstance(issue, str):
        return issue
    if not isinstance(issue, dict):
        return str(issue)
    section = issue.get("section")
    text = (
        issue.get("issue")
        or issue.get("gap")
        or issue.get("why_blocking")
        or issue.get("suggestion")
    )
    if section and text:
        return f"{section}: {text}"
    return text or json.dumps(issue, ensure_ascii=False)


def extract_remaining_issues(state_path):
    if not state_path or not Path(state_path).exists():
        return []
    try:
        state = json.loads(Path(state_path).read_text(encoding="utf-8-sig"))
    except (json.JSONDecodeError, OSError):
        return []

    explicit_issues = state.get("remaining_blocking_issues")
    if isinstance(explicit_issues, list) and explicit_issues:
        return [summarize_issue(issue) for issue in explicit_issues[:10]]

    client_readiness = state.get("client_readiness")
    if isinstance(client_readiness, dict):
        client_issues = client_readiness.get("blocking_issues")
        if isinstance(client_issues, list) and client_issues:
            return [summarize_issue(issue) for issue in client_issues[:10]]

    iterations = state.get("iterations", [])
    if not iterations:
        return []
    last = iterations[-1]
    blocking_count = sum(
        r.get("blocking_count", 0)
        for r in last.get("reviews", {}).values()
        if isinstance(r, dict)
    )
    if blocking_count == 0:
        return []
    summary = []
    for reviewer, data in last.get("reviews", {}).items():
        if isinstance(data, dict) and data.get("blocking_count", 0) > 0:
            summary.append(f"{reviewer}: {data['blocking_count']} blocking issue(s)")
    return summary


# ---------------------------------------------------------------------------
# Markdown parser — line-by-line, simple state machine
# ---------------------------------------------------------------------------

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
HORIZONTAL_RULE_RE = re.compile(r"^-{3,}\s*$")
TABLE_LINE_RE = re.compile(r"^\s*\|.*\|\s*$")


def split_table_row(line):
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def is_table_separator(line):
    if not TABLE_LINE_RE.match(line):
        return False
    cells = split_table_row(line)
    if not cells:
        return False
    return all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in cells)


def add_markdown_table(doc, rows):
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=col_count)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    for row_index, row_cells in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(col_count):
            cell_text = row_cells[col_index] if col_index < len(row_cells) else ""
            paragraph = cells[col_index].paragraphs[0]
            # Tables use Arial 12pt body but with no first-line indent to
            # keep cell content readable.
            _apply_std_paragraph_format(paragraph, first_line_indent=Cm(0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            render_inline(paragraph, cell_text, bold_outer=(row_index == 0))

    doc.add_paragraph()


def render_markdown(doc, md_text):
    """Stream through lines, classify each, emit docx paragraphs."""
    lines = md_text.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()

        if not line.strip():
            index += 1
            continue

        # Markdown table (header line + separator + data rows)
        if (
            TABLE_LINE_RE.match(line)
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1])
        ):
            table_rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and TABLE_LINE_RE.match(lines[index].rstrip()):
                table_rows.append(split_table_row(lines[index].rstrip()))
                index += 1
            add_markdown_table(doc, table_rows)
            continue

        # Horizontal rule — render as a blank paragraph (visual break)
        if HORIZONTAL_RULE_RE.match(line):
            doc.add_paragraph()
            index += 1
            continue

        # Heading (#, ##, ###, ####) — all render the same: bold body
        # paragraph with no indent. Numbering is part of the text.
        m = HEADING_RE.match(line)
        if m:
            add_heading_paragraph(doc, m.group(2).strip())
            index += 1
            continue

        # Blockquote (> ...) — italic 11pt with left indent
        m = BLOCKQUOTE_RE.match(line)
        if m:
            add_blockquote_paragraph(doc, m.group(1).strip())
            index += 1
            continue

        # Bullet (- * +)
        m = BULLET_RE.match(line)
        if m:
            add_list_item(doc, m.group(1).strip(), ordered=False)
            index += 1
            continue

        # Numbered (1. 2. 3.)
        m = NUMBERED_RE.match(line)
        if m:
            add_list_item(doc, m.group(1).strip(), ordered=True)
            index += 1
            continue

        # Plain paragraph
        add_body_paragraph(doc, line)
        index += 1


# ---------------------------------------------------------------------------
# Inline formatting — bold/italic/code recognition inside one paragraph
# ---------------------------------------------------------------------------

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
CODE_RE = re.compile(r"`([^`]+)`")


def render_inline(paragraph, text, *, bold_outer=False, italic_outer=False, size=None):
    """Add inline-formatted runs to `paragraph`.

    `bold_outer` / `italic_outer` / `size` apply to all plain-text segments
    (the wrapping paragraph type's defaults). Inline ** / * / ` then layer
    on top of those defaults.
    """

    def emit(segment_text, *, segment_bold=False, segment_italic=False, segment_code=False):
        if not segment_text:
            return
        run = paragraph.add_run(segment_text)
        run.bold = bold_outer or segment_bold
        run.italic = italic_outer or segment_italic
        run.font.name = "Consolas" if segment_code else FONT_NAME
        run.font.size = size if size is not None else FONT_SIZE_BODY
        if not segment_code:
            # apply east-asian font preservation too
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), FONT_NAME)
            rfonts.set(qn("w:hAnsi"), FONT_NAME)
            rfonts.set(qn("w:cs"), FONT_NAME)
            rfonts.set(qn("w:eastAsia"), FONT_NAME)

    cursor = 0
    while cursor < len(text):
        match = None
        kind = None
        for k, regex in (("bold", BOLD_RE), ("italic", ITALIC_RE), ("code", CODE_RE)):
            m = regex.search(text, cursor)
            if m and (match is None or m.start() < match.start()):
                match = m
                kind = k
        if match is None:
            emit(text[cursor:])
            break
        if match.start() > cursor:
            emit(text[cursor:match.start()])
        emit(
            match.group(1),
            segment_bold=(kind == "bold"),
            segment_italic=(kind == "italic"),
            segment_code=(kind == "code"),
        )
        cursor = match.end()


# ---------------------------------------------------------------------------
# Language-specific typography (currently a no-op)
# ---------------------------------------------------------------------------

def apply_locale_typography(md_text, language):
    """Reserved hook for future locale-specific typography substitutions.

    The plugin is English-only as of 0.0.35 and this function is a no-op;
    it remains here to keep the CLI shape stable for downstream callers.
    """
    return md_text


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Convert legal memo markdown to docx.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--template-id", default="classical-memo")
    parser.add_argument("--final-status", default=None)
    parser.add_argument("--state", default=None)
    parser.add_argument(
        "--language",
        default="en",
        choices=("en",),
        help="Memo language. Plugin is English-only as of 0.0.35; flag kept for CLI stability.",
    )
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        sys.stderr.write(f"ERROR: input markdown not found: {input_path}\n")
        sys.exit(1)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    md_text = input_path.read_text(encoding="utf-8")
    md_text = apply_locale_typography(md_text, args.language)

    doc = Document()
    apply_page_setup(doc)
    configure_default_style(doc)

    fallback_banners = extract_fallback_banners(args.state)
    needs_banner = (args.final_status and not args.final_status.startswith("approved")) or bool(fallback_banners)
    if needs_banner:
        remaining = extract_remaining_issues(args.state)
        add_warning_banner(doc, args.final_status, remaining, fallback_banners=fallback_banners)

    render_markdown(doc, md_text)

    doc.save(str(output_path))
    sys.stdout.write(f"OK: wrote {output_path}\n")


if __name__ == "__main__":
    main()
