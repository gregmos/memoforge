"""Smoke tests for md_to_docx.py fallback_banners rendering.

Verifies that:
- state.json.fallback_banners[] gets rendered as bullets in the warning banner.
- Banner fires even when final_status starts with "approved" if fallback_banners is non-empty.
- Missing/empty fallback_banners doesn't break rendering for the standard non-approved path.

These tests require python-docx (the same dep the script itself uses). They
write a real .docx to a temp dir, then parse it back to check the banner text.
"""
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

PLUGIN_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = PLUGIN_ROOT / "lib" / "docx-render" / "scripts" / "md_to_docx.py"

try:
    from docx import Document  # type: ignore
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def write_minimal_state(workdir: Path, **fields) -> Path:
    state = {
        "task_id": "test",
        "fallback_banners": [],
        "remaining_blocking_issues": [],
        **fields,
    }
    p = workdir / "state.json"
    p.write_text(json.dumps(state), encoding="utf-8")
    return p


def write_minimal_md(workdir: Path) -> Path:
    p = workdir / "in.md"
    p.write_text("# Title\n\nBody paragraph.\n", encoding="utf-8")
    return p


def run_renderer(input_md: Path, output_docx: Path, state: Path, final_status: str | None) -> int:
    args = [
        sys.executable, str(SCRIPT),
        "--input", str(input_md),
        "--output", str(output_docx),
        "--state", str(state),
    ]
    if final_status:
        args += ["--final-status", final_status]
    result = subprocess.run(args, capture_output=True, text=True)
    if result.returncode != 0:
        print("STDERR:", result.stderr, file=sys.stderr)
    return result.returncode


def extract_all_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


@unittest.skipUnless(HAS_DOCX, "python-docx not installed")
class MdToDocxBannerTests(unittest.TestCase):

    def test_fallback_banners_rendered_as_bullets(self):
        with tempfile.TemporaryDirectory() as td:
            wd = Path(td)
            md = write_minimal_md(wd)
            out = wd / "out.docx"
            state = write_minimal_state(
                wd,
                fallback_banners=[
                    "Research summary mode — full IRAC analysis not performed per user choice.",
                    "MCP rate-limited; some sources fetched via WebFetch.",
                ],
            )
            code = run_renderer(md, out, state, final_status="fallback_research_summary_delivered")
            self.assertEqual(code, 0)
            self.assertTrue(out.exists())
            text = extract_all_text(out)
            self.assertIn("Research summary mode", text)
            self.assertIn("MCP rate-limited", text)
            self.assertIn("Pipeline fallbacks that fired during this run", text)
            self.assertIn("RESEARCH SUMMARY MODE", text)  # custom title for the research-summary path

    def test_no_banner_for_approved_with_no_fallbacks(self):
        with tempfile.TemporaryDirectory() as td:
            wd = Path(td)
            md = write_minimal_md(wd)
            out = wd / "out.docx"
            state = write_minimal_state(wd)
            code = run_renderer(md, out, state, final_status="approved_on_v2")
            self.assertEqual(code, 0)
            text = extract_all_text(out)
            self.assertNotIn("MANUAL REVIEW REQUIRED", text)
            self.assertNotIn("Pipeline fallbacks", text)

    def test_approved_status_still_shows_banner_when_fallbacks_present(self):
        # Edge case: user took summary path, mediator marked it "approved" but
        # fallback_banners contains a banner. Banner must still render.
        with tempfile.TemporaryDirectory() as td:
            wd = Path(td)
            md = write_minimal_md(wd)
            out = wd / "out.docx"
            state = write_minimal_state(
                wd,
                fallback_banners=["MCP rate-limited fallback fired during research."],
            )
            code = run_renderer(md, out, state, final_status="approved_on_v1")
            self.assertEqual(code, 0)
            text = extract_all_text(out)
            self.assertIn("MCP rate-limited fallback", text)

    def test_forced_exit_with_remaining_and_fallbacks(self):
        with tempfile.TemporaryDirectory() as td:
            wd = Path(td)
            md = write_minimal_md(wd)
            out = wd / "out.docx"
            state = write_minimal_state(
                wd,
                fallback_banners=["MCP unavailable; some sources fetched from canonical portals."],
                remaining_blocking_issues=["Issue 1: missing case status."],
            )
            code = run_renderer(md, out, state, final_status="forced_exit_on_v3_with_remaining_issues")
            self.assertEqual(code, 0)
            text = extract_all_text(out)
            self.assertIn("REVIEWER NOTES NOT FULLY RESOLVED", text)
            self.assertIn("Pipeline fallbacks", text)
            self.assertIn("MCP unavailable", text)
            self.assertIn("Remaining blocking issues", text)
            self.assertIn("missing case status", text)


if __name__ == "__main__":
    unittest.main()
